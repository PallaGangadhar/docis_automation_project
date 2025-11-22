import requests
import secrets
import datetime
from flask import request
from db import db_connection, get_tab_info, get_resgression_id_by_log_id,get_resgression_id
from docx.shared import RGBColor, Pt

per_page = 5

def get_query_string():
    query_string = request.url.split('?')[1]
    return query_string

def send_req(*args, **kwargs):
    if kwargs:
        message = kwargs.get('output')
        random_string = kwargs.get('random_string')
        rerun = 'False'
        tab_info = get_tab_info(random_string)
        rerun = 'True' if tab_info is None else 'False'
        tab_info = f"{tab_info}_{random_string}"
        requests.post("http://localhost:5000/send_message", json={"data": message, 'query_string':tab_info, 'rerun':rerun},headers = {'Content-type': 'application/json'})
    elif args:
        requests.post("http://localhost:5000/send_message", json={"data": args[0]},headers = {'Content-type': 'application/json'})
        # requests.post("http://localhost:5000/send_message", json={"data": args[0], 'query_string': args[1]},headers = {'Content-type': 'application/json'})


def generate_key():
    ''' genearate key '''
    return secrets.token_hex(6)


def genearate_list_of_dict(r_dates,type_dates,graph_data):
    new_list=[]
    for r_date in r_dates:
        if r_date not in type_dates:
            graph_data.append({'date':r_date,'count':0}) 
    graph_data.sort(key = lambda x:x['date'],reverse=True)
    [new_list.append(i) for i in graph_data]
    return new_list


def convert_str_to_date(date,d_format):
    datetime_str = datetime.datetime.strptime(date, d_format)
    return datetime_str

def convert_date_to_str(date,d_format):
    return date.strftime(d_format)

def filter_regression(cmts_type,search_reg,query):
    params = []
    # Apply Filters
    if cmts_type:
        query += " AND cmts_type = %s"
        params.append(cmts_type)

    if search_reg:
        query += " AND LOWER(regression_name) LIKE LOWER(%s)"
        params.append(f"%{search_reg}%")
    return query, params

def filter_testcase(tc_name,module_type,device_type,query):
    params = []
    # Apply Filters
    if tc_name:
        query += " AND LOWER(testcase_name) LIKE LOWER(%s)"
        params.append(f"%{tc_name}%")

    if device_type:
        query += " AND devices_details.device_id = %s"
        params.append(device_type)

    if module_type:
        query += " AND modules_details.modules_id = %s"
        params.append(module_type)
    return query, params

def filter_devices(search_device,query):
    params = []
    # Apply Filters
    if search_device:
        query += " AND LOWER(device_name) LIKE LOWER(%s)"
        params.append(f"%{search_device}%")
    return query, params

def filter_modules(search_module, device_type,query):
    params = []
    # Apply Filters
    if device_type:
        query += " AND devices_details.device_id= %s"
        params.append(device_type)

    if search_module:
        query += " AND LOWER(module_name) LIKE LOWER(%s)"
        params.append(f"%{search_module}%")
    return query, params

def get_offset_details(**kwargs):
    curr, conn = db_connection()
    # Base Query
    table_name = kwargs.get('table_name')
    offset =  kwargs.get('offset')
    per_page = kwargs.get('per_page')
    filter_for = kwargs.get('filter_for')
    order_by_field = None
    params = None
    order = 'ASC'
    if table_name:
        query = f"SELECT * FROM {table_name} WHERE 1=1"

    if table_name == "regression":
        cmts_type = kwargs.get('cmts_type')
        search_reg = kwargs.get('search_reg')
        query, params = filter_regression(cmts_type, search_reg, query)
        order_by_field = "date_added"
        order = 'DESC'
    
    elif table_name == "testcase_details":
        tc_name = kwargs.get('tc_name')
        module_type = kwargs.get('module_type')
        device_type = kwargs.get('device_type')
        query = f"SELECT {table_name}.*,modules_details.module_name,devices_details.device_name\
        FROM {table_name},modules_details,devices_details where {table_name}.modules_id = \
        modules_details.modules_id and modules_details.device_id=devices_details.device_id\
         and 1=1"
        query, params = filter_testcase(tc_name,module_type,device_type,query)
        order_by_field = "testcase_id"

    elif table_name == "devices_details":
        search_device = kwargs.get('search_devices')
        query, params = filter_devices(search_device, query)
        order_by_field = "device_id"
    
    elif table_name == "modules_details":
        search_module = kwargs.get('search_module')
        device_type = kwargs.get('device_type')
        query = "SELECT devices_details.device_name,modules_id, module_name FROM public.modules_details, \
                devices_details where devices_details.device_id=modules_details.device_id and 1=1"
        query, params = filter_modules(search_module, device_type, query)
        order_by_field = "modules_id"
    
    elif filter_for == "get_tc_execution_status":
        testcase_id = kwargs.get('testcase_id')
        print("testcase id === ", testcase_id)
        query = f"""SELECT DISTINCT 
            testcase_number, 
            testcase_name,
            testcase_id
        FROM 
            testcase_details where 
            testcase_id IN {testcase_id} """
        # query = f"""SELECT DISTINCT 
        #     tc.testcase_number, 
        #     tc.testcase_name, 
        #     reg.status 
        # FROM 
        #     testcase_details AS tc
        # LEFT JOIN 
        #     regression_logs_details AS reg 
        #     ON tc.testcase_id = reg.testcase_id
        # WHERE 
        #     tc.testcase_id IN {testcase_id} """
        
        # query = """
        #     SELECT  
        #         tc.testcase_number, 
        #         tc.testcase_name, 
        #         reg.status,
        #         reg.regression_id
        #     FROM 
        #         testcase_details AS tc
        #     LEFT JOIN 
        #         regression_logs_details AS reg 
        #         ON tc.testcase_id = reg.testcase_id
        #     LEFT JOIN 
        #         regression AS r 
        #         ON reg.regression_id = r.regression_id
        #     WHERE 
        #         r.regression_string = %s"""
    

    if table_name:
        query += f" ORDER BY {order_by_field} {order} LIMIT %s OFFSET %s"
        params.extend([per_page, offset])
        curr.execute(query, tuple(params))
    else:
        params = []
        query += f"LIMIT %s OFFSET %s"
        params.extend([per_page, offset])
        curr.execute(query, params)
    offset_details = curr.fetchall()
    curr.close()
    conn.close()
    return offset_details

# Function to count total regression details with filters
def get_table_total_details(**kwargs):
    curr, conn = db_connection()
    table_name = kwargs.get('table_name')
    filter_for = kwargs.get('filter_for')

    # Base Query
    query = f"SELECT * FROM {table_name} WHERE 1=1"
    params = None
    if table_name == "regression":
        cmts_type = kwargs.get('cmts_type')
        search_reg = kwargs.get('search_reg')
        query, params = filter_regression(cmts_type, search_reg,query)
    
    elif table_name == "testcase_details":
        tc_name = kwargs.get('tc_name')
        module_type = kwargs.get('module_type')
        device_type = kwargs.get('device_type')
        query = f"SELECT {table_name}.*,modules_details.module_name,devices_details.device_name\
        FROM {table_name},modules_details,devices_details where {table_name}.modules_id = \
        modules_details.modules_id and modules_details.device_id=devices_details.device_id"
        query, params = filter_testcase(tc_name, module_type, device_type, query)
        
    elif table_name == "devices_details":
        search_device = kwargs.get('search_devices')
        query, params = filter_devices(search_device, query)

    elif table_name == "modules_details":
        search_module = kwargs.get('search_module')
        device_type = kwargs.get('device_type')
        query = "SELECT devices_details.device_name,modules_id, module_name FROM public.modules_details, \
                devices_details where devices_details.device_id=modules_details.device_id "
        query, params = filter_modules(search_module, device_type, query)
    
    elif filter_for == "get_tc_execution_status":
        testcase_id = kwargs.get('testcase_id')
        random_string = kwargs.get('random_string')
       
        query = f"""SELECT DISTINCT
            testcase_number, 
            testcase_name,
            testcase_id

        FROM 
            testcase_details

        WHERE 
            testcase_id in {testcase_id} """
        # query = """
            # SELECT  
            #     tc.testcase_number, 
            #     tc.testcase_name, 
            #     reg.status,
            #     reg.regression_id
            # FROM 
            #     testcase_details AS tc
            # LEFT JOIN 
            #     regression_logs_details AS reg 
            #     ON tc.testcase_id = reg.testcase_id
            # LEFT JOIN 
            #     regression AS r 
            #     ON reg.regression_id = r.regression_id
            # WHERE 
            #     r.regression_string = %s"""
    
    if params:
        curr.execute(query, tuple(params))
    else:
        # params = [random_string]

        curr.execute(query)
    total = len(curr.fetchall())
    curr.close()
    conn.close()
    return total


def get_date_diff(from_date, to_date):
    date_format = "%Y-%m-%d"
    from_date_dt = datetime.datetime.strptime(from_date, date_format)
    to_date_dt = datetime.datetime.strptime(to_date, date_format)
    # Calculate the difference
    date_difference = (to_date_dt - from_date_dt).days
    if date_difference > 31 and date_difference <= 364: 
        return ">30 and <=364"
    elif date_difference >=365:
        return ">365"
    return False


def table_summary(status,fail_in,execution_time,tc_logs_path,tc_id, log_id, query_string):
    print("Query String === ", query_string)
    if log_id != None:
        reg_id = get_resgression_id_by_log_id(log_id)
    else:
        reg_id = get_resgression_id(query_string)
    r=requests.post("http://localhost:5000/add_regression_logs", json={"tc_id":tc_id, "status":status,'fail_in':fail_in,'execution_time':execution_time,'tc_logs_path':tc_logs_path, "r_id":reg_id, "log_id": log_id},headers = {'Content-type': 'application/json'})
    print(r.status_code)


def validate_tc_status(current_status, existing_status):
    return True if current_status == existing_status else False
    

# def find_replace_docx(doc, output_file, data):
#     # Replace in paragraphs
#     for key in data:
#         find_text = f"{key['testcase_name']}-TBT"
#         replace_text = f"{key['status']}"
#         print("Find Text: ", find_text)
#         print("Replace Text: ", replace_text)
#         color = RGBColor(0, 255, 0) if  replace_text == "PASS" else RGBColor(255, 0, 0)  # Green color (R, G, B)
#         for para in doc.paragraphs:
#             if find_text in para.text:
#                 for run in para.runs:
#                     if find_text in run.text:
                        
#                         run.text = run.text.replace(find_text, replace_text)

#                         # Apply formatting
#                         run.font.bold = True                     # make text bold
#                         run.font.name = 'Arial'                # font name
#                         run.font.size = Pt(10)                   # font size in points
#                         run.font.color.rgb = color  # green color (R, G, B)

#         # Replace in tables
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         if find_text in para.text:
#                             for run in para.runs:
#                                 if find_text in run.text:
#                                     print("find text === ", find_text)
#                                     run.text = run.text.replace(find_text, replace_text)

#                                     # Apply formatting
#                                     run.font.bold = True
#                                     run.font.name = 'Arial'
#                                     run.font.size = Pt(10)
#                                     run.font.color.rgb = RGBColor(0, 255, 0)

#     doc.save(output_file)


def find_replace_docx(doc, output_file, data):

    def replace_in_paragraph(paragraph, find_text, replace_text):
        full_text = ''.join(run.text for run in paragraph.runs)
        if full_text in find_text and "TBT" in full_text:
            new_text = full_text.replace(find_text, replace_text)
            for run in paragraph.runs:
                run.text = ''

            # Create a new run for replaced text
            new_run = paragraph.add_run(new_text)
            new_run.font.bold = True
            new_run.font.name = 'Arial'
            new_run.font.size = Pt(10)
            new_run.font.color.rgb = RGBColor(0, 255, 0) if replace_text == "PASS" else RGBColor(255, 0, 0)

    # Replace in paragraphs
    for key in data:
        find_text = f"{key['testcase_name']}-TBT"
        replace_text = f"{key['status']}"
        for para in doc.paragraphs:
            replace_in_paragraph(para, find_text, replace_text)

    # Replace in tables
    for key in data:
        find_text = f"{key['testcase_name']}-TBT"
        replace_text = f"{key['status']}"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, find_text, replace_text)

    doc.save(output_file)
# Example usage
# find_replace_docx("docs\\new.docx", "docs\\test_update2.docx", "This is to verify the NTP status (Ipv4)-TBT", "PASS")